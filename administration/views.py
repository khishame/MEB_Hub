from django.shortcuts import render,redirect
from login.models import Campus,Admin,Student
from bus.models import ScheduleCode,Bus_schedule,Bus,Bus_Stats
from .models import Admin_Action
from django.http import HttpResponseRedirect,HttpResponse
from datetime import datetime,timedelta
from datetime import date
from django.forms import Form
from events.models import Event
from django.utils import timezone
from django.db.models.functions import TruncDate
from django.db.models import Count
from django.utils.timezone import now, timedelta
from reportlab.pdfgen import canvas
from docx import Document
import csv
from io import BytesIO


# Create your views here.
def admin_home(request):
  admin = Admin.objects.all().get(admin_id=request.session['admin_id'])
  today = timezone.localdate()
  actions = Admin_Action.objects.annotate(action_date=TruncDate('datetime')).filter(admin_id=admin.admin_id,action_date=today)
  events = Event.objects.filter(date__year=today.year, date__month=today.month)
  cnt_routes=ScheduleCode.objects.count()
  context={
    "admin":admin,
    "actions" :actions,
    "events" :events,
    "initials": f"{admin.name[0].upper()}{admin.surname[0].upper()}",
    "cnt_routes":cnt_routes
  }
  return render(request,"admin/admin.html",context
  )

#campus functionalilities

#bus functionalities
def bus_menu(request):
  initials=request.session.get('initials')
  return render(request,"admin/buses/bus_menu.html",{'initials':initials})

def add_all_campuses_view(request):
  list = ScheduleCode.objects.all()
  initials = request.session.get("initials")
  return render(request,"admin/buses/add_all_camp.html",{
    "list": list,"initials":initials
  })
  
def add_bus_schedule(request,code):
  
  schedule_c = ScheduleCode.objects.all().get(schedule_code=code)
  bus_list = Bus.objects.all()
  initials = request.session.get("initials")
  
  errors = {}
  
  if request.method == 'POST':

    bus = Bus.objects.all().get(bus_id=int(request.POST['bus']))
    start_time = request.POST['start_time']
    last_time = request.POST['last_time']
    duration = int(request.POST['duration'])
    
    # Basic field validation
    if not bus or not start_time or not last_time or not duration:
        errors['form'] = "All fields are required."
    
    s_time = datetime.strptime(start_time, '%H:%M')
    l_time = datetime.strptime(last_time, '%H:%M')
    
    if l_time < s_time:
      errors['time error'] = "Last time should not be earlier than start time"
      
    # if duration > 2 or duration < 1 :
    #   errors['duration error'] = "Duration should be between 1 and 2"
    
    
    
    n_time = datetime.strptime(start_time, '%H:%M')
      
    if not errors:
      Bus_schedule.objects.filter(schedule_code=code).delete()
      n_time = datetime.strptime(start_time, '%H:%M')
      while n_time < l_time:
        n_time = s_time + timedelta(minutes=duration)
        
        bus_s = Bus_schedule(departure=schedule_c.campus1,destination=schedule_c.campus2,
                               departure_time=s_time.time(),arrival_time=n_time.time(),duration=duration,
                               bus_id=bus,schedule_code=schedule_c)
          
        bus_s.save()
          
        s_time = n_time + timedelta(minutes=duration)
        n_time = s_time + timedelta(minutes=duration)
        
        if n_time > l_time:
          break
    
        bus_s = Bus_schedule(departure=schedule_c.campus2,destination=schedule_c.campus1,
                          departure_time=s_time.time(),arrival_time=n_time.time(),duration=duration,
                          bus_id=bus,schedule_code=schedule_c)
        bus_s.save()
          
        s_time = n_time + timedelta(minutes=duration)
          
      admin = Admin.objects.all().get(admin_id=request.session['admin_id'])
      addAction(admin_id=admin,record_type="Generated bus schedule",icon="bi bi-bus-front")
      return HttpResponseRedirect(redirect_to='/administration/bus_menu')
    else:
        return render(request,"admin/buses/add_schedule.html",{
          "schedule": schedule_c,
          "bus_list": bus_list,"initials":initials,
          "errors": errors
        })
  
  return render(request,"admin/buses/add_schedule.html",{
        "schedule": schedule_c,
        "bus_list": bus_list,"initials":initials,
        "errors": errors
      })
  
def events_menu(request):
  Event.objects.filter(date__lt=date.today()).delete()  # automatically delete events
  admin_id = request.session.get('admin_id')
  initials = request.session.get("initials")

  if admin_id is None:
    return redirect('admin_home')  # or handle expired session

  admin = Admin.objects.select_related('campus_id').get(admin_id=admin_id)
  return render(request,"admin/events/events_menu.html",{'admin':admin,'initials':initials})

def addAction(admin_id: int,record_type: str,icon: str):
  action = Admin_Action(action_type=record_type,admin_id=admin_id, icon=icon,
                        datetime = timezone.now())
  action.save()
  

def view_all_actions(request):
  admin = Admin.objects.all().get(admin_id=request.session['admin_id'])
  today = timezone.now().date()
  actions = Admin_Action.objects.annotate(action_date=TruncDate('datetime')).filter(admin_id=admin.admin_id,
                                                                                    action_date=today)
  events = Event.objects.filter(date__year=today.year, date__month=today.month)
  cnt_routes = ScheduleCode.objects.count()
  context = {
    "admin": admin,
    "actions": actions,
    "events": events,
    "initials": f"{admin.name[0].upper()}{admin.surname[0].upper()}",
    "cnt_routes": cnt_routes
  }
  
  return render(request,"admin/admin.html",context)

def add_bus(request):
  if request.method == 'POST':
    form = Form(request.POST)
    if form.is_valid():
      bus_name = request.POST["bus_name"]
      campus_id = request.POST["campus"]

      campus = Campus.objects.all().get(campus_id=campus_id)
      bus = Bus(bus_name=bus_name,campus_id=campus)

      bus.save()
      
      admin = Admin.objects.all().get(admin_id=request.session['admin_id'])


      
      addAction(admin_id=admin,record_type="Added a new bus",icon="bi bi-bus-front")
      
      return redirect('view_all_actions')
    
    else:
       return  render(request,"admin/buses/add_bus_html",{
         "form": form
       })

  campus_list = Campus.objects.all()
  initials = request.session.get("initials")
  return render(request,"admin/buses/add_bus.html",{
    "campus_list": campus_list,"initials":initials
  })


def remove_bus(request):
   if request.method == 'POST':
      form = Form(request.POST)
      initials = request.session.get("initials")

      if form.is_valid:
        bus_id = request.POST["bus"]

        bus = Bus.objects.all().get(bus_id=bus_id)
        
        bus.delete()
        
        admin = Admin.objects.all().get(admin_id=request.session['admin_id'])


        addAction(admin_id=admin,record_type="Bus has been removed",icon="bi bi-x-circle")
        
        return redirect('view_all_actions')
      else:
        return render(request,"admin/buses/remove_bus.html",{
          "form": form,"initials":initials
        })
   initials = request.session.get("initials")

   bus_list = Bus.objects.all()
   return render(request,"admin/buses/remove_bus.html",{
     "bus_list": bus_list,"initials":initials
   })

def bus_schedule_stats(request):
  total_views = Bus_Stats.objects.count()

  # Views per schedule
  views_per_schedule = Bus_Stats.objects.values(
    'schedule_code__schedule_code',
    'schedule_code__campus1',
    'schedule_code__campus2'
  ).annotate(view_count=Count('id')).order_by('-view_count')

  # Views by day
  views_by_day = Bus_Stats.objects.annotate(
    date=TruncDate('viewed_at')
  ).values('date').annotate(count=Count('id')).order_by('date')

  # Views in the last 7 days
  seven_days_ago = timezone.now() - timedelta(days=7)
  recent_views = Bus_Stats.objects.filter(viewed_at__gte=seven_days_ago).annotate(
    date=TruncDate('viewed_at')
  ).values('date').annotate(count=Count('id')).order_by('date')

  # Prepare data for charts
  schedule_labels = [
    f"{item['schedule_code__schedule_code']} - {item['schedule_code__campus1']} and {item['schedule_code__campus2']}"
    for item in views_per_schedule
  ]
  schedule_counts = [item['view_count'] for item in views_per_schedule]

  day_labels = [item['date'] for item in views_by_day]
  day_counts = [item['count'] for item in views_by_day]

  recent_labels = [item['date'] for item in recent_views]
  recent_counts = [item['count'] for item in recent_views]

  context = {
    'total_views': total_views,
    'views_per_schedule': list(zip(schedule_labels, schedule_counts)),
    'views_by_day': list(zip(day_labels, day_counts)),
    'recent_views': list(zip(recent_labels, recent_counts)),
    'most_viewed': views_per_schedule[0] if views_per_schedule else None,
  }

  return render(request, 'admin/buses/schedule_stats.html', context)

def user_management(request):
    initials=request.session.get("initials")
    return render(request,'admin/user_management.html',{"initials":initials})

def analytics(request):
    initials=request.session.get("initials")
    return render(request,'admin/analytics.html',{'initials':initials})
      
  
def bus_schedule_stats(request):
  total_views = Bus_Stats.objects.count()

  # Views per schedule
  views_per_schedule = Bus_Stats.objects.values(
    'schedule_code__schedule_code',
    'schedule_code__campus1',
    'schedule_code__campus2'
  ).annotate(view_count=Count('id')).order_by('-view_count')

  # Views by day
  views_by_day = Bus_Stats.objects.annotate(
    date=TruncDate('viewed_at')
  ).values('date').annotate(count=Count('id')).order_by('date')

  # Views in the last 7 days
  seven_days_ago = timezone.now() - timedelta(days=7)
  recent_views = Bus_Stats.objects.filter(viewed_at__gte=seven_days_ago).annotate(
    date=TruncDate('viewed_at')
  ).values('date').annotate(count=Count('id')).order_by('date')

  # Prepare data for charts
  schedule_labels = [
    f"{item['schedule_code__schedule_code']} - {item['schedule_code__campus1']} and {item['schedule_code__campus2']}"
    for item in views_per_schedule
  ]
  schedule_counts = [item['view_count'] for item in views_per_schedule]

  day_labels = [item['date'] for item in views_by_day]
  day_counts = [item['count'] for item in views_by_day]

  recent_labels = [item['date'] for item in recent_views]
  recent_counts = [item['count'] for item in recent_views]

  initials = request.session.get("initials")


  context = {
    'total_views': total_views,
    'views_per_schedule': list(zip(schedule_labels, schedule_counts)),
    'views_by_day': list(zip(day_labels, day_counts)),
    'recent_views': list(zip(recent_labels, recent_counts)),
    'most_viewed': views_per_schedule[0] if views_per_schedule else None,
    'initials':initials
  }

  return render(request, 'admin/buses/schedule_stats.html', context)


def student_report(request):
    students = Student.objects.all()

    total_students = students.count()
    bus_users = Bus_Stats.objects.count()
    non_bus_users = total_students - bus_users

    campus_counts = students.values('campus_id__campus_name').annotate(count=Count('studentNumber')).order_by('-count')

    latest_login = students.order_by('-login_time').first()

    initials=request.session.get("initials")

    context = {
      'students': students,
      'total_students': total_students,
      'bus_users': bus_users,
      'non_bus_users': non_bus_users,
      'campus_counts': campus_counts,
      'latest_login': latest_login,
      'initials': initials
    }
    return render(request, 'admin/student_report.html', context)


def student_report_pdf(request):
  response = HttpResponse(content_type='application/pdf')
  response['Content-Disposition'] = 'attachment; filename="student_report.pdf"'

  p = canvas.Canvas(response)
  p.setFont("Helvetica", 12)

  # Metrics
  total_students = Student.objects.count()
  bus_users = Bus_Stats.objects.count()
  non_bus_users = Bus_Stats.objects.count()
  campus_counts = Student.objects.values('campus_id__campus_name').annotate(count=Count('studentNumber'))

  y = 800
  p.drawString(100, y, "Student Report")
  y -= 30

  p.drawString(50, y, f"Total Students: {total_students}")
  y -= 20
  p.drawString(50, y, f"Students Using Bus: {bus_users}")
  y -= 20
  p.drawString(50, y, f"Students Not Using Bus: {non_bus_users}")
  y -= 30

  p.drawString(50, y, "Students per Campus:")
  y -= 20
  for campus in campus_counts:
    p.drawString(70, y, f"{campus['campus_id__campus_name']}: {campus['count']}")
    y -= 20
    if y < 100:
      p.showPage()
      y = 800

  y -= 10
  p.drawString(50, y, "Detailed Student List:")
  y -= 30

  students = Student.objects.select_related('campus_id').all()
  for student in students:
    line = f"{student.studentNumber} - {student.name} {student.surname} - {student.studentEmail} - {student.campus_id.campus_name} - Bus: {'Yes' if Bus_Stats.viewed_at else 'No'}"
    p.drawString(50, y, line)
    y -= 20
    if y < 100:
      p.showPage()
      y = 800

  p.save()
  return response


def student_report_docx(request):
  doc = Document()
  doc.add_heading('Student Report', 0)

  # Metrics
  total_students = Student.objects.count()
  bus_users = Bus_Stats.objects.count()
  non_bus_users = Bus_Stats.objects.count()
  campus_counts = Student.objects.values('campus_id__campus_name').annotate(count=Count('studentNumber'))

  # Summary Section
  doc.add_paragraph(f'Total Students: {total_students}')
  doc.add_paragraph(f'Students Using Bus: {bus_users}')
  doc.add_paragraph(f'Students Not Using Bus: {non_bus_users}')
  doc.add_paragraph("Students per Campus:")
  for campus in campus_counts:
    doc.add_paragraph(f"• {campus['campus_id__campus_name']}: {campus['count']}", style='List Bullet')

  # Table of students
  doc.add_paragraph('Student Details:')
  table = doc.add_table(rows=1, cols=5)
  table.style = 'Table Grid'
  hdr_cells = table.rows[0].cells
  hdr_cells[0].text = 'Student No.'
  hdr_cells[1].text = 'Name'
  hdr_cells[2].text = 'Surname'
  hdr_cells[3].text = 'Email'
  hdr_cells[4].text = 'Campus'

  students = Student.objects.select_related('campus_id').all()
  for student in students:
    row_cells = table.add_row().cells
    row_cells[0].text = student.studentNumber
    row_cells[1].text = student.name
    row_cells[2].text = student.surname
    row_cells[3].text = student.studentEmail
    row_cells[4].text = student.campus_id.campus_name if student.campus_id else "N/A"

  # Response
  response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
  response['Content-Disposition'] = 'attachment; filename=student_report.docx'
  doc.save(response)
  return response


def student_report_csv(request):
  # Create HTTP response with CSV content type
  response = HttpResponse(content_type='text/csv')
  response['Content-Disposition'] = 'attachment; filename="student_report.csv"'

  writer = csv.writer(response)

  # Write report title
  writer.writerow(['Student Report'])

  # Write column headers
  writer.writerow(['Student Number', 'Name', 'Surname', 'Email', 'Campus', 'Uses Bus'])

  students = Student.objects.select_related('campus_id').all()

  for student in students:
    writer.writerow([
      student.studentNumber,
      student.name,
      student.surname,
      student.studentEmail,
      student.campus_id.campus_name if student.campus_id else 'N/A',
      'Yes' if Bus_Stats.objects.count() else 'No'
    ])

  return response


def full_report(request):
  one_week_ago = now() - timedelta(days=7)
  admins = Admin.objects.all()
  events = Event.objects.all()
  buses = Bus.objects.all()
  students = Student.objects.all()
  schedules = Bus_schedule.objects.all()
  recent_admins = admins.filter(created_at__gte=one_week_ago) if hasattr(Admin, 'created_at') else []
  recent_events = events.filter(created_at__gte=one_week_ago) if hasattr(Event, 'created_at') else []
  recent_schedules = schedules.filter(departure_time__gte=now().time())
  recent_actions = Admin_Action.objects.filter(datetime__gte=one_week_ago)
  initials = request.session.get("initials")

  # Import your Student model accordingly

  return render(request, 'admin/full_report.html', {
    'admin_count': admins.count(),
    'event_count': events.count(),
    'bus_count': buses.count(),
    'schedule_count': schedules.count(),
    'recent_admins': recent_admins,
    'recent_events': recent_events,
    'recent_schedules': recent_schedules,
    'recent_actions': recent_actions,
    'total_students': students.count(),
    'initials': initials    # Added student count here
  })


def full_report_pdf(request):
  one_week_ago = now() - timedelta(days=7)
  admins = Admin.objects.all()
  events = Event.objects.all()
  buses = Bus.objects.all()
  schedules = Bus_schedule.objects.all()
  recent_admins = admins.filter(created_at__gte=one_week_ago) if hasattr(Admin, 'created_at') else []
  recent_events = events.filter(created_at__gte=one_week_ago) if hasattr(Event, 'created_at') else []
  recent_schedules = schedules.filter(departure_time__gte=now().time())
  recent_actions = Admin_Action.objects.filter(datetime__gte=one_week_ago)

  response = HttpResponse(content_type='application/pdf')
  response['Content-Disposition'] = 'attachment; filename="full_report.pdf"'

  buffer = BytesIO()
  p = canvas.Canvas(buffer)
  y = 800

  def write_line(text, indent=100):
    nonlocal y
    if y < 50:
      p.showPage()
      y = 800
    p.drawString(indent, y, text)
    y -= 20

  write_line(f"System Report ({now().date()})")
  write_line(
    f"Admins: {admins.count()}, Events: {events.count()}, Buses: {buses.count()}, Schedules: {schedules.count()}")
  write_line("")

  write_line("Recent Admins:")
  for admin in recent_admins:
    write_line(f"{admin.name} {admin.surname} - {admin.email}", indent=120)

  write_line("Recent Events:")
  for event in recent_events:
    write_line(f"{event.description} at {event.location} on {event.date}", indent=120)

  write_line("Recent Schedules:")
  for sched in recent_schedules:
    write_line(f"{sched.departure} to {sched.destination} ({sched.departure_time}-{sched.arrival_time})", indent=120)

  write_line("Admin Actions:")
  for action in recent_actions:
    write_line(f"{action.admin_id.name}: {action.action_type} at {action.datetime}", indent=120)

  p.showPage()
  p.save()

  pdf = buffer.getvalue()
  buffer.close()
  response.write(pdf)
  return response


def full_report_docx(request):
  one_week_ago = now() - timedelta(days=7)
  admins = Admin.objects.all()
  events = Event.objects.all()
  buses = Bus.objects.all()
  schedules = Bus_schedule.objects.all()
  recent_admins = admins.filter(created_at__gte=one_week_ago) if hasattr(Admin, 'created_at') else []
  recent_events = events.filter(created_at__gte=one_week_ago) if hasattr(Event, 'created_at') else []
  recent_schedules = schedules.filter(departure_time__gte=now().time())
  recent_actions = Admin_Action.objects.filter(datetime__gte=one_week_ago)

  doc = Document()
  doc.add_heading('System Report', 0)
  doc.add_paragraph(f"Date: {now().date()}")
  doc.add_paragraph(
    f"Admins: {admins.count()}, Events: {events.count()}, Buses: {buses.count()}, Schedules: {schedules.count()}")

  doc.add_heading("Recent Admins", level=1)
  for admin in recent_admins:
    doc.add_paragraph(f"{admin.name} {admin.surname} - {admin.email}")

  doc.add_heading("Recent Events", level=1)
  for event in recent_events:
    doc.add_paragraph(f"{event.description} at {event.location} on {event.date}")

  doc.add_heading("Recent Schedules", level=1)
  for sched in recent_schedules:
    doc.add_paragraph(f"{sched.departure} to {sched.destination} ({sched.departure_time}-{sched.arrival_time})")

  doc.add_heading("Admin Actions", level=1)
  for action in recent_actions:
    doc.add_paragraph(f"{action.admin_id.name}: {action.action_type} at {action.datetime}")

  buffer = BytesIO()
  doc.save(buffer)
  buffer.seek(0)

  response = HttpResponse(buffer.getvalue(),
                          content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
  response['Content-Disposition'] = 'attachment; filename="full_report.docx"'
  return response


def full_report_csv(request):
  one_week_ago = now() - timedelta(days=7)
  admins = Admin.objects.all()
  events = Event.objects.all()
  buses = Bus.objects.all()
  schedules = Bus_schedule.objects.all()

  recent_admins = admins.filter(created_at__gte=one_week_ago) if hasattr(Admin, 'created_at') else []
  recent_events = events.filter(created_at__gte=one_week_ago) if hasattr(Event, 'created_at') else []
  recent_schedules = schedules.filter(departure_time__gte=now().time())
  recent_actions = Admin_Action.objects.filter(datetime__gte=one_week_ago)

  response = HttpResponse(content_type='text/csv')
  response['Content-Disposition'] = 'attachment; filename="full_report.csv"'
  writer = csv.writer(response)

  writer.writerow(["Section", "Details"])
  writer.writerow(["Counts",
                   f"Admins: {admins.count()}, Events: {events.count()}, Buses: {buses.count()}, Schedules: {schedules.count()}"])

  writer.writerow(["Recent Admins"])
  for admin in recent_admins:
    writer.writerow(["", f"{admin.name} {admin.surname} - {admin.email}"])

  writer.writerow(["Recent Events"])
  for event in recent_events:
    writer.writerow(["", f"{event.description} at {event.location} on {event.date}"])

  writer.writerow(["Recent Schedules"])
  for sched in recent_schedules:
    writer.writerow(["", f"{sched.departure} to {sched.destination} ({sched.departure_time}-{sched.arrival_time})"])

  writer.writerow(["Admin Actions"])
  for action in recent_actions:
    writer.writerow(["", f"{action.admin_id.name}: {action.action_type} at {action.datetime}"])

  return response